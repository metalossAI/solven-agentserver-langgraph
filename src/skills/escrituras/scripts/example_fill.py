#!/usr/bin/env python3
"""
Example fast fill script for a DOCX notary template using {{PLACEHOLDER}} keys.

Replace this with a model-specific script (e.g. 42_modelo_fill.py) after analysing
your template with scripts/analyze_docx_placeholders.py.

Usage:
  python scripts/example_fill.py template.docx data.json -o filled.docx
  python scripts/example_fill.py template.docx '{"NOMBRE":"Juan"}' -o filled.docx

Requires: python-docx (pip install python-docx)
"""

import re
import sys
import json
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("python-docx required: pip install python-docx", file=sys.stderr)
    sys.exit(1)

PLACEHOLDER_PATTERN = re.compile(r"\{\{([^}]+)\}\}")


def replace_placeholders_in_paragraph(paragraph, data: dict) -> None:
    """Replace {{KEY}} in a paragraph in place, keeping formatting where possible."""
    full = paragraph.text
    if "{{" not in full:
        return
    for key, value in data.items():
        placeholder = "{{" + key + "}}"
        if placeholder in full:
            full = full.replace(placeholder, str(value))
    if full != paragraph.text:
        # Clear and set new text (simple approach; can be refined with runs)
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = full
        else:
            paragraph.add_run(full)


def fill_docx(template_path: str, data: dict, output_path: str) -> None:
    """Load DOCX, replace all {{KEY}} with data values, save to output_path."""
    doc = Document(template_path)
    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, data)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholders_in_paragraph(paragraph, data)
    doc.save(output_path)


def main():
    if len(sys.argv) < 4 or "-o" not in sys.argv:
        print(
            "Usage: example_fill.py <template.docx> <data.json|'{\"K\":\"v\"}'> -o <output.docx>",
            file=sys.stderr,
        )
        sys.exit(1)
    template_path = sys.argv[1]
    data_arg = sys.argv[2]
    out_idx = sys.argv.index("-o")
    output_path = sys.argv[out_idx + 1]

    path = Path(data_arg)
    if path.suffix.lower() == ".json" and path.exists():
        with open(path, encoding="utf-8") as f:
            data = json.load(f)
    else:
        data = json.loads(data_arg)

    fill_docx(template_path, data, output_path)
    print("Written:", output_path)


if __name__ == "__main__":
    main()
